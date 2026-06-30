"""Renderer Word del motor de export (python-docx). Espejo de _pdf: mapea `datos` por estructura.

escalares → "Resumen" + párrafos; analisis/contexto_datos → texto; lista de dicts → tabla;
lista de escalares → bullets; dict simple → key/value. `_sheets` → una sección por hoja.
"""
import io
from typing import Any, Dict


def build_docx(nombre: str, datos: Dict[str, Any]) -> bytes:
    """Genera un .docx a partir de (nombre, datos). datos asume primitivos (sin coerción de tipos)."""
    from docx import Document

    doc = Document()
    doc.add_heading(nombre, level=0)

    sheets = datos.get("_sheets")
    if isinstance(sheets, dict):
        for i, (sheet_name, sheet_datos) in enumerate(sheets.items()):
            if i > 0:
                doc.add_page_break()
            doc.add_heading(sheet_name, level=1)
            _escribir_seccion(doc, sheet_datos)
    else:
        _escribir_seccion(doc, datos)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _escribir_seccion(doc: Any, datos: Dict[str, Any]) -> None:
    """Vuelca un dict de datos al documento, espejando el mapeo del renderer PDF."""
    # ── Escalares ────────────────────────────────────────────────────────────
    scalars = {k: v for k, v in datos.items()
               if not isinstance(v, (list, dict)) and k not in ("titulo",)}
    if scalars:
        doc.add_heading("Resumen", level=2)
        for key, val in scalars.items():
            p = doc.add_paragraph()
            p.add_run(f"{key.replace('_', ' ').capitalize()}: ").bold = True
            p.add_run(str(val))

    # ── Texto largo ──────────────────────────────────────────────────────────
    for key in ("analisis", "contexto_datos"):
        if isinstance(datos.get(key), str):
            doc.add_heading(key.replace("_", " ").capitalize(), level=2)
            for line in datos[key].split("\n"):
                doc.add_paragraph(line)

    # ── Tablas (listas) ──────────────────────────────────────────────────────
    for key, val in datos.items():
        if not isinstance(val, list) or not val:
            continue
        doc.add_heading(key.replace("_", " ").capitalize(), level=2)
        if isinstance(val[0], dict):
            headers = list(val[0].keys())
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Light Grid Accent 1"
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h.replace("_", " ").capitalize()
            for item in val:
                cells = table.add_row().cells
                for i, h in enumerate(headers):
                    cells[i].text = str(item.get(h, ""))
        else:
            for item in val:
                doc.add_paragraph(str(item), style="List Bullet")

    # ── Dicts simples (excluye claves privadas) ──────────────────────────────
    for key, val in datos.items():
        if key.startswith("_") or not isinstance(val, dict):
            continue
        doc.add_heading(key.replace("_", " ").capitalize(), level=2)
        for k, v in val.items():
            p = doc.add_paragraph()
            p.add_run(f"{k}: ").bold = True
            p.add_run(str(v))
